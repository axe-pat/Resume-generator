#!/usr/bin/env python3
"""Build the visibly labeled profile-maxing reference resume.

Design system:
- Base preset: compact_reference_guide.
- First-page pattern: simplified memo_masthead (name, role, contact line).
- Named override: resume_one_page_reference.
  Letter portrait; 0.50 in top/bottom, 0.55 in left/right;
  Arial 9 pt body; compact exact line spacing; no tables.
- Persistent safety furniture: red counterfactual warning in header and footer.

This script is intentionally scoped to the isolated profile_maxing_lab and never
writes to the factual resume pipeline.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "resume" / "maxed_profile_reference_resume.docx"

FONT = "Arial"
INK = "17202A"
MUTED = "505A64"
ACCENT = "8A1C1C"
RULE = "9EA7B0"
CONTENT_WIDTH_IN = 8.5 - (0.55 * 2)


SUMMARY = (
    "USC Marshall MBA candidate and current AI product intern with five years of "
    "backend engineering across mobility, data infrastructure, fintech, and healthcare. "
    "Builds decision infrastructure that turns complex systems and fragmented workflows "
    "into trustworthy user action."
)


EDUCATION = [
    {
        "name": "University of Southern California, Marshall School of Business",
        "location": "Los Angeles, CA",
        "date": "May 2027",
        "detail": "Master of Business Administration (STEM); Dean's Merit Scholarship",
    },
    {
        "name": "Thapar Institute of Engineering and Technology",
        "location": "Patiala, India",
        "date": "Aug 2020",
        "detail": "Bachelor of Engineering, Computer Engineering; Merit Scholarship; Springer-published AI research (90+ citations)",
    },
]


EXPERIENCE = [
    {
        "company": "FlairX",
        "context": "AI recruiting and interview workflow startup",
        "date": "Current; start date TBD",
        "role": "Product Manager Intern (AI Products) - provisional title; confirm against official record",
        "bullets": [
            "Turned an enterprise refusal to outsource final interviews into a 0-to-1 internal-panel product; led a two-week build of privacy-safe M365 scheduling and evidence-linked AI scoring, cutting recruiter overhead 42% and unblocking Genpact.",
            "Owned FlairX's Ceipal ATS integration and the pull-only MVP decision after the vendor API blocked write-back; removed 80% of double entry, retained a flagship account, and opened a marketplace acquisition channel.",
            "Converted a 20-minute avatar-vendor cap into a build-vs-buy decision and swappable rendering layer; negotiated cost from ~$0.33 to $0.10/min, cutting vendor cost ~70% and removing a single point of enterprise failure.",
        ],
    },
    {
        "company": "Gojek",
        "context": "Southeast Asian ride-hailing marketplace",
        "date": "Jan 2025 - Jul 2025",
        "role": "Senior Software Engineer | counterfactual lens: Product Engineer, Supply & Marketplace",
        "bullets": [
            "Found a p95 3.8-second fare-quote tail hidden behind a healthy average; reframed latency as a duopoly conversion problem, cut tail latency 70%, and enabled ~28K incremental monthly rides.",
            "Redesigned the external-fleet contract around four-hour capacity commitments and confidence-weighted matching, increasing active supply 18% and reducing pickup ETA 1.5 minutes across launch corridors.",
        ],
    },
    {
        "company": "Hevo Data",
        "context": "Enterprise ELT data platform",
        "date": "Nov 2023 - Jan 2025",
        "role": "Software Engineer II | counterfactual lens: Product Engineer, Enterprise Data",
        "bullets": [
            "Proved enterprise buyers valued auditability over sub-minute freshness and shaped a batch-first transactional roadmap; onboarded eight enterprise customers in 90 days while migrating the base without churn.",
            "Built an evidence-backed AI incident layer that grouped 40-60 cascading alerts into one root-cause card, cutting diagnosis from 45 minutes to under five and mean time to recovery 40%.",
        ],
    },
    {
        "company": "Intuit",
        "context": "QuickBooks monetization and billing systems",
        "date": "Aug 2022 - Oct 2023",
        "role": "Software Engineer II | counterfactual lens: Product Engineer, Billing Platform",
        "bullets": [
            "Exposed a 50K-account integrity backlog hidden across 12 teams and launched a governed reconciliation overlay that auto-resolved 3K+ discrepancies monthly and cut manual corrections 68%.",
        ],
    },
    {
        "company": "Optum",
        "context": "Healthcare data and care-network platform",
        "date": "Jul 2020 - Aug 2022",
        "role": "Software Engineer | counterfactual lens: Product Engineer, Care Platform",
        "bullets": [
            "Productized provider onboarding through a reusable 80/20 transformation contract and joint clinical certification path, cutting integration time from six months to 10 weeks.",
        ],
    },
]


PROJECT = {
    "company": "Recruiting Engine",
    "context": "Product Manager & AI-Native Builder",
    "date": "Mar 2026 - Present",
    "bullets": [
        "Built and operated an AI recruiting decision engine across 151 commits and 542 release tests, unifying 2,514 roles, 560 organizations, 846 contacts, and 849 touchpoints into evidence-backed application and relationship actions.",
    ],
}


SKILLS = [
    ("Product", "Discovery, product strategy, roadmap sequencing, experimentation, pricing, platform/API products, enterprise workflows, responsible AI"),
    ("Technical", "Python, SQL, TypeScript, REST APIs, distributed systems, data pipelines, LLM evaluation, observability"),
    ("Operating", "PRDs, decision memos, metric design, Figma, Jira, stakeholder alignment, incident leadership"),
]


def set_cellless_font(run, *, size: float, bold: bool = False, italic: bool = False, color: str = INK) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if value and keep is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not value and keep is not None:
        p_pr.remove(keep)


def set_bottom_border(paragraph, color: str = RULE, size: int = 5, space: int = 1) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_bullet_numbering(document: Document) -> int:
    """Create a real Word bullet numbering definition and return numId."""
    numbering = document.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "360")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    p_pr.append(ind)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def configure_document() -> tuple[Document, int]:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.50)
    section.bottom_margin = Inches(0.50)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.20)
    section.footer_distance = Inches(0.20)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(11.2)

    styles = doc.styles
    if "Resume Section" not in styles:
        style = styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Resume Section"]
    style.base_style = styles["Heading 1"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(10.3)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(ACCENT)
    style.paragraph_format.space_before = Pt(5)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    style.paragraph_format.line_spacing = Pt(11.3)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    hp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    hp.paragraph_format.line_spacing = Pt(8)
    set_shading(hp, "FDECEC")
    hr = hp.add_run("COUNTERFACTUAL REFERENCE | INVENTED DETAILS | NOT FOR EXTERNAL USE")
    set_cellless_font(hr, size=7.2, bold=True, color=ACCENT)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fp.paragraph_format.line_spacing = Pt(8)
    fr = fp.add_run("REFERENCE ONLY - VERIFY AND REWRITE EVERY CLAIM BEFORE PUBLIC USE")
    set_cellless_font(fr, size=7.2, bold=True, color=ACCENT)

    return doc, add_bullet_numbering(doc)


def section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Resume Section")
    p.add_run(text.upper())
    set_bottom_border(p, color=RULE, size=4, space=1)


def name_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(17.5)
    r = p.add_run("AKSHAT PATHAK")
    set_cellless_font(r, size=16.8, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(10)
    r = p.add_run("TECHNICAL PRODUCT MANAGER / PRODUCT ENGINEER")
    set_cellless_font(r, size=9.4, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(9.5)
    r = p.add_run(
        "Los Angeles, CA | Akshat.Pathak.2027@marshall.usc.edu | "
        "(213) 325-0609 | linkedin.com/in/akshats-pathak"
    )
    set_cellless_font(r, size=8.7, color=MUTED)


def add_summary(doc: Document) -> None:
    section_heading(doc, "Product Management")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(SUMMARY)
    set_cellless_font(r, size=9.35)


def add_header_line(doc: Document, left_bold: str, left_normal: str, date: str, *, before: float = 0.7) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(10.4)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT)
    set_keep_with_next(p)
    r = p.add_run(left_bold)
    set_cellless_font(r, size=9.65, bold=True)
    r = p.add_run(f" - {left_normal}")
    set_cellless_font(r, size=9.0, color=MUTED)
    r = p.add_run(f"\t{date}")
    set_cellless_font(r, size=9.0, bold=True)


def add_role_line(doc: Document, role: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(9.9)
    set_keep_with_next(p)
    r = p.add_run(role)
    set_cellless_font(r, size=8.85, italic=True, color=MUTED)


def add_bullet(doc: Document, num_id: int, text: str) -> None:
    p = doc.add_paragraph()
    set_num(p, num_id)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.125)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(10.7)
    r = p.add_run(text)
    set_cellless_font(r, size=9.1)


def add_education(doc: Document) -> None:
    section_heading(doc, "Education")
    for idx, item in enumerate(EDUCATION):
        add_header_line(doc, item["name"], item["location"], item["date"], before=0.4 if idx else 0)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(9.7)
        r = p.add_run(item["detail"])
        set_cellless_font(r, size=8.8, italic=True, color=MUTED)


def add_experience(doc: Document, num_id: int) -> None:
    section_heading(doc, "Experience")
    for idx, item in enumerate(EXPERIENCE):
        add_header_line(doc, item["company"], item["context"], item["date"], before=1.6 if idx else 0)
        add_role_line(doc, item["role"])
        for bullet in item["bullets"]:
            add_bullet(doc, num_id, bullet)


def add_project(doc: Document, num_id: int) -> None:
    section_heading(doc, "Selected Product Build")
    add_header_line(doc, PROJECT["company"], PROJECT["context"], PROJECT["date"], before=0)
    for bullet in PROJECT["bullets"]:
        add_bullet(doc, num_id, bullet)


def add_skills(doc: Document) -> None:
    section_heading(doc, "Skills")
    for label, value in SKILLS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(10.1)
        r = p.add_run(f"{label}: ")
        set_cellless_font(r, size=8.75, bold=True)
        r = p.add_run(value)
        set_cellless_font(r, size=8.75)


def build() -> Path:
    doc, num_id = configure_document()
    name_block(doc)
    add_summary(doc)
    add_education(doc)
    add_experience(doc, num_id)
    add_project(doc, num_id)
    add_skills(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    try:
        print(build())
    except Exception as exc:  # pragma: no cover - CLI safety
        print(f"build failed: {exc}", file=sys.stderr)
        raise
