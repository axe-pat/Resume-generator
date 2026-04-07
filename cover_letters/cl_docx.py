#!/usr/bin/env python3
"""
cl_docx.py — Generate a clean, paste-ready cover letter .docx
==============================================================
Produces a minimal formatted document:

  [Company] Product Team,

  [paragraph 1]

  [paragraph 2]

  [paragraph 3]

  [paragraph 4]

  Best,
  Akshat

Usage (standalone):
  python cl_docx.py <cl_body.txt> <output.docx> <Company>

Usage (from Python):
  from cl_docx import generate_cl_docx
  generate_cl_docx(cl_body_raw, output_path, company="Stripe")

cl_body_raw: the CL text BEFORE add_salutation_signoff() is called.
             The function strips any trailing signature block automatically.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────────────────────
FONT_NAME     = "Times New Roman"
FONT_SIZE_PT  = 11
LINE_SPACING  = Pt(14)      # ~1.15 line spacing for 11pt
SPACE_AFTER   = Pt(10)      # gap between paragraphs
MARGIN_IN     = Inches(1.0) # 1" all sides


# ─────────────────────────────────────────────────────────────────────────────
# Body extractor — strips trailing "Akshat Pathak / email" signature
# ─────────────────────────────────────────────────────────────────────────────
def _extract_body_paragraphs(cl_body_raw: str) -> list[str]:
    """
    Given raw CL text (no salutation, but may have a trailing signature block),
    return a list of body paragraph strings — no signature, no blank elements.
    """
    lines = cl_body_raw.strip().splitlines()

    # Find and drop the trailing signature (Akshat Pathak line + email)
    sig_idx = None
    for i in range(len(lines) - 1, -1, -1):
        if "akshat pathak" in lines[i].lower():
            sig_idx = i
            break
    if sig_idx is not None:
        # Also drop any blank lines immediately before the signature
        while sig_idx > 0 and not lines[sig_idx - 1].strip():
            sig_idx -= 1
        lines = lines[:sig_idx]

    # Group remaining lines into paragraphs (split on blank lines)
    paragraphs: list[str] = []
    current: list[str] = []
    for line in lines:
        if line.strip():
            current.append(line.strip())
        else:
            if current:
                paragraphs.append(" ".join(current))
                current = []
    if current:
        paragraphs.append(" ".join(current))

    return [p for p in paragraphs if p]


# ─────────────────────────────────────────────────────────────────────────────
# Docx builder
# ─────────────────────────────────────────────────────────────────────────────
def _set_run_font(run, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)
    run.bold      = bold
    # Force font via XML for compatibility (Word sometimes ignores Python-level font)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),    FONT_NAME)
    rFonts.set(qn("w:hAnsi"),    FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.insert(0, rFonts)


def _add_paragraph(doc: Document, text: str, space_after: Pt = SPACE_AFTER,
                   bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = space_after
    p.paragraph_format.line_spacing  = LINE_SPACING
    # Remove any default style indentation
    p.paragraph_format.left_indent   = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    if text:
        run = p.add_run(text)
        _set_run_font(run, bold=bold)


def generate_cl_docx(
    cl_body_raw: str,
    output_path:  "str | Path",
    company:      str = "",
) -> Path:
    """
    Generate a clean CL .docx.

    Args:
        cl_body_raw:  CL text before salutation/signoff is added (raw Step 2
                      output with <!-- --> stripped). Trailing signature block
                      is stripped automatically.
        output_path:  Where to write the .docx file.
        company:      Company name for the salutation line.

    Returns:
        Path to the written .docx file.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Extract body paragraphs (no signature)
    body_paras = _extract_body_paragraphs(cl_body_raw)

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = MARGIN_IN
        section.bottom_margin = MARGIN_IN
        section.left_margin   = MARGIN_IN
        section.right_margin  = MARGIN_IN

    # ── Default style: strip any inherited spacing ────────────────────────────
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(FONT_SIZE_PT)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

    # Remove the empty paragraph Document() adds by default
    if doc.paragraphs and not doc.paragraphs[0].text:
        p_elem = doc.paragraphs[0]._element
        p_elem.getparent().remove(p_elem)

    # ── Salutation ───────────────────────────────────────────────────────────
    company_clean = company.strip() if company.strip() else "Hiring"
    salutation    = f"{company_clean} Product Team,"
    _add_paragraph(doc, salutation, space_after=Pt(14))

    # ── Body paragraphs ──────────────────────────────────────────────────────
    for para_text in body_paras:
        _add_paragraph(doc, para_text, space_after=Pt(14))

    # ── Sign-off ─────────────────────────────────────────────────────────────
    _add_paragraph(doc, "Best,",   space_after=Pt(2))
    _add_paragraph(doc, "Akshat",  space_after=Pt(0))

    doc.save(str(output_path))
    return output_path


# ─────────────────────────────────────────────────────────────────────────────
# Standalone CLI
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python cl_docx.py <cl_body.txt> <output.docx> [Company]")
        sys.exit(1)

    input_txt   = Path(sys.argv[1])
    output_docx = Path(sys.argv[2])
    company_arg = sys.argv[3] if len(sys.argv) > 3 else ""

    if not input_txt.exists():
        print(f"[ERROR] Input file not found: {input_txt}")
        sys.exit(1)

    body_text = input_txt.read_text(encoding="utf-8")
    out = generate_cl_docx(body_text, output_docx, company=company_arg)
    print(f"OK: {out}")
