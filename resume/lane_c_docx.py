#!/usr/bin/env python3
"""Build compact Lane C resume and cover-letter DOCX files from JSON payloads."""

from __future__ import annotations

import json
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


FONT = "Times New Roman"


def _paragraph_by_text(doc: Document, text: str):
    for paragraph in doc.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise ValueError(f"Template paragraph not found: {text!r}")


def _clear_run_content(run_element):
    """Keep a template run's properties while removing its visible content."""
    for child in list(run_element):
        if child.tag != qn("w:rPr"):
            run_element.remove(child)


def _append_run_text(run_element, text: str):
    """Append text and real Word tab elements to an OOXML run."""
    parts = text.split("\t")
    for index, part in enumerate(parts):
        if part:
            node = OxmlElement("w:t")
            if part[:1].isspace() or part[-1:].isspace():
                node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            node.text = part
            run_element.append(node)
        if index < len(parts) - 1:
            run_element.append(OxmlElement("w:tab"))


def _clone_paragraph(template_paragraph, run_specs: list[tuple[str, int]] | None = None):
    """Clone one authority paragraph, optionally replacing its runs in-place."""
    paragraph_element = deepcopy(template_paragraph._p)
    if run_specs is None:
        return paragraph_element

    for child in list(paragraph_element):
        if child.tag != qn("w:pPr"):
            paragraph_element.remove(child)

    source_runs = template_paragraph.runs
    if not source_runs:
        raise ValueError(f"Template paragraph has no runs: {template_paragraph.text!r}")
    for text, run_index in run_specs:
        source = source_runs[min(run_index, len(source_runs) - 1)]._r
        cloned_run = deepcopy(source)
        _clear_run_content(cloned_run)
        _append_run_text(cloned_run, text)
        paragraph_element.append(cloned_run)
    return paragraph_element


def _split_label_detail(text: str) -> tuple[str, str]:
    """Bold the entity name while keeping descriptors and locations regular."""
    for separator in (" — ", " – "):
        if separator in text:
            label, detail = text.split(separator, 1)
            return label, separator + detail
    if " (" in text:
        label, detail = text.split(" (", 1)
        return label, " (" + detail
    return text, ""


def _build_resume_from_template(data: dict, output: Path):
    """Clone the curated Amazon DOCX and replace body content using its paragraph patterns."""
    template_path = Path(data["template_path"]).expanduser().resolve()
    if not template_path.exists():
        raise FileNotFoundError(template_path)
    if template_path == output:
        raise ValueError("Template and output paths must differ; older artifacts are never overwritten")

    output.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(template_path, output)
    doc = Document(output)

    # These prototypes come from the manually curated Amazon authority. Cloning
    # their OOXML keeps rules, tabs, spacing, bullets, justification, and run styling.
    name_proto = doc.paragraphs[0]
    contact_proto = doc.paragraphs[1]
    section_proto = _paragraph_by_text(doc, "PRODUCT MANAGEMENT")
    summary_proto = doc.paragraphs[3]
    education_header_proto = _paragraph_by_text(doc, "University of Southern California, Marshall School of Business – Los Angeles, CA\tMay 2027")
    degree_proto = _paragraph_by_text(doc, "Master of Business Administration (STEM)")
    education_bullet_proto = _paragraph_by_text(doc, "Honors: Dean’s Merit Scholarship")
    spacer_proto = doc.paragraphs[9]
    entry_header_proto = _paragraph_by_text(doc, "FlairX AI (AI-first interview and hiring platform) – San Francisco, CA\tJun 2026 – Aug 2026")
    role_proto = _paragraph_by_text(doc, "AI Product Manager Intern")
    experience_bullet_proto = doc.paragraphs[17]
    skill_proto = doc.paragraphs[40]

    new_paragraphs = [
        _clone_paragraph(name_proto, [(data["name"], 0)]),
        _clone_paragraph(contact_proto, [(data["contact"], 0)]),
    ]

    profile = data.get("profile")
    if profile:
        profile_header = data.get("summary_section_header", "PROFILE")
        new_paragraphs.append(_clone_paragraph(section_proto, [(profile_header.upper(), 0)]))
        new_paragraphs.append(_clone_paragraph(summary_proto, [(profile, 0)]))

    new_paragraphs.append(_clone_paragraph(section_proto, [("EDUCATION", 0)]))
    for index, row in enumerate(data.get("education", [])):
        if index:
            new_paragraphs.append(_clone_paragraph(spacer_proto))
        school_label, school_detail = _split_label_detail(row["school"])
        header_runs = [(school_label, 0)]
        if school_detail:
            header_runs.append((school_detail, 1))
        if row.get("date"):
            header_runs.append(("\t" + row["date"], 2))
        new_paragraphs.append(_clone_paragraph(education_header_proto, header_runs))
        new_paragraphs.append(_clone_paragraph(degree_proto, [(row["degree"], 0)]))
        for bullet in row.get("bullets", []):
            new_paragraphs.append(_clone_paragraph(education_bullet_proto, [(bullet, 0)]))

    for section in data.get("sections", []):
        new_paragraphs.append(_clone_paragraph(section_proto, [(section["title"].upper(), 0)]))
        for index, entry in enumerate(section.get("entries", [])):
            if index:
                new_paragraphs.append(_clone_paragraph(spacer_proto))
            organization_label, organization_detail = _split_label_detail(entry["organization"])
            header_runs = [(organization_label, 0)]
            if organization_detail:
                header_runs.append((organization_detail, 2))
            if entry.get("date"):
                header_runs.append(("\t" + entry["date"], 3))
            new_paragraphs.append(_clone_paragraph(entry_header_proto, header_runs))
            if entry.get("role"):
                new_paragraphs.append(_clone_paragraph(role_proto, [(entry["role"], 0)]))
            for bullet in entry.get("bullets", []):
                new_paragraphs.append(_clone_paragraph(experience_bullet_proto, [(bullet, 0)]))

    if data.get("skills"):
        skills_header = data.get("skills_section_header", "SKILLS & ADDITIONAL")
        new_paragraphs.append(_clone_paragraph(section_proto, [(skills_header.upper(), 0)]))
        for row in data["skills"]:
            if isinstance(row, dict):
                label, value = row["bold_label"], row["text"]
            else:
                label, value = row
            new_paragraphs.append(
                _clone_paragraph(skill_proto, [(label + ":", 0), (" " + value, 1)])
            )

    body = doc._element.body
    section_properties = body.sectPr
    if section_properties is None:
        raise ValueError("Template is missing section properties")
    for child in list(body):
        if child is not section_properties:
            body.remove(child)
    for paragraph_element in new_paragraphs:
        body.insert(len(body) - 1, paragraph_element)

    doc.save(output)


def _set_font(run, size: float = 10, *, bold: bool = False, italic: bool = False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def _set_cellless_defaults(doc: Document, size: float = 10):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    style.font.size = Pt(size)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0


def _add_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def _section_header(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    _add_bottom_border(p)
    _set_font(p.add_run(text.upper()), 10, bold=True)
    return p


def _header_line(doc: Document, left: str, right: str = "", *, size: float = 10):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    _set_font(p.add_run(left), size, bold=True)
    if right:
        _set_font(p.add_run("\t" + right), size)
    return p


def _subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    _set_font(p.add_run(text), 10, italic=True)
    return p


def _bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    _set_font(p.add_run(text), 10)
    return p


def build_resume(data: dict, output: Path):
    if data.get("template_path"):
        _build_resume_from_template(data, output)
        return

    doc = Document()
    _set_cellless_defaults(doc, 10)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.48)
    sec.bottom_margin = Inches(0.48)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    _set_font(p.add_run(data["name"]), 16, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _set_font(p.add_run(data["contact"]), 9.5)

    if data.get("profile"):
        _section_header(doc, "Profile")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _set_font(p.add_run(data["profile"]), 10)

    _section_header(doc, "Education")
    for index, row in enumerate(data.get("education", [])):
        if index:
            doc.add_paragraph().paragraph_format.space_after = Pt(0)
        _header_line(doc, row["school"], row.get("date", ""))
        _subtitle(doc, row["degree"])
        for bullet in row.get("bullets", []):
            _bullet(doc, bullet)

    for section in data.get("sections", []):
        _section_header(doc, section["title"])
        for index, entry in enumerate(section.get("entries", [])):
            if index:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.line_spacing = 0.2
                spacer.paragraph_format.space_after = Pt(0)
                _set_font(spacer.add_run(""), 2)
            _header_line(doc, entry["organization"], entry.get("date", ""))
            if entry.get("role"):
                _subtitle(doc, entry["role"])
            for bullet in entry.get("bullets", []):
                _bullet(doc, bullet)

    if data.get("skills"):
        _section_header(doc, "Skills & Additional")
        for label, value in data["skills"]:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.space_after = Pt(0)
            _set_font(p.add_run(label + ": "), 10, bold=True)
            _set_font(p.add_run(value), 10)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def build_cover_letter(data: dict, output: Path):
    doc = Document()
    _set_cellless_defaults(doc, 11)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    _set_font(p.add_run(data["name"]), 16, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    _set_font(p.add_run(data["contact"]), 10)

    for line in [data["date"], *data["address"]]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _set_font(p.add_run(line), 11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(8)
    _set_font(p.add_run(data["salutation"]), 11)

    for text in data["paragraphs"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.08
        _set_font(p.add_run(text), 11)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    _set_font(p.add_run(data.get("closing", "Sincerely,")), 11)
    p = doc.add_paragraph()
    _set_font(p.add_run(data["name"]), 11, bold=True)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main():
    if len(sys.argv) != 2:
        raise SystemExit("Usage: lane_c_docx.py <payload.json>")
    payload_path = Path(sys.argv[1]).resolve()
    data = json.loads(payload_path.read_text(encoding="utf-8"))
    output = Path(data["output_path"]).expanduser().resolve()
    if data["kind"] == "resume":
        build_resume(data, output)
    elif data["kind"] == "cover_letter":
        build_cover_letter(data, output)
    else:
        raise SystemExit(f"Unknown kind: {data['kind']}")
    print(output)


if __name__ == "__main__":
    main()
